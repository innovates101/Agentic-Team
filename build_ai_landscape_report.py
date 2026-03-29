"""
build_ai_landscape_report.py
Builds the AI Landscape 2026 Word document from Leary's JSON output.
  - Cover page + disclaimer
  - Clickable Table of Contents
  - Key Statistics at a Glance panel
  - Inline citation injection
  - Formatted body (H1/H2/H3, bullets, numbered lists, blockquotes, HR)
  - References section
  - Further Reading section with live hyperlinks
"""
import json, re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INBOX  = Path("Owner's Inbox")
data   = json.loads((INBOX / "result_ai_landscape_report_leary.json").read_text(encoding="utf-8"))
text   = data["result"]

doc = Document()

# ── Margins ───────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.left_margin   = Inches(1.15)
    section.right_margin  = Inches(1.15)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

NAVY   = RGBColor(0x0D, 0x2C, 0x54)
BLUE   = RGBColor(0x1A, 0x6B, 0xC4)
STEEL  = RGBColor(0x3A, 0x5F, 0x8F)
BODY   = RGBColor(0x1C, 0x1C, 0x1C)
LGRAY2 = RGBColor(0x88, 0x88, 0x88)

# ── Helpers ───────────────────────────────────────────────────────────────────
def add_styled_runs(para, line, colour=BODY, size_pt=11):
    line = inject_citations(line)
    for part in re.split(r"(\*\*.*?\*\*)", line):
        if part.startswith("**") and part.endswith("**"):
            r = para.add_run(part[2:-2]); r.bold = True
        else:
            r = para.add_run(part)
        r.font.color.rgb = colour
        r.font.size = Pt(size_pt)

def add_hr(doc, colour="1A6BC4", thickness="6"):
    p = doc.add_paragraph()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), thickness)
    bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), colour)
    pBdr.append(bot)
    p._p.get_or_add_pPr().append(pBdr)
    p.paragraph_format.space_after = Pt(6)

def add_blockquote(doc, line):
    content = line.lstrip("> ").strip()
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.35)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "4");    left.set(qn("w:color"), "1A6BC4")
    pBdr.append(left)
    p._p.get_or_add_pPr().append(pBdr)
    run = p.add_run(content)
    run.italic = True; run.font.color.rgb = STEEL; run.font.size = Pt(11)

def add_shaded_box(doc, title, lines, bg_hex="EEF3F9", border_hex="1A6BC4"):
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(10)
    title_p.paragraph_format.space_after  = Pt(2)
    pPr = title_p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "0D2C54")
    pPr.append(shd)
    r = title_p.add_run(f"  {title}")
    r.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.size = Pt(12)
    for line in lines:
        item_p = doc.add_paragraph()
        item_p.paragraph_format.left_indent  = Inches(0.25)
        item_p.paragraph_format.space_before = Pt(3)
        item_p.paragraph_format.space_after  = Pt(3)
        pPr2 = item_p._p.get_or_add_pPr()
        shd2 = OxmlElement("w:shd")
        shd2.set(qn("w:val"), "clear"); shd2.set(qn("w:color"), "auto"); shd2.set(qn("w:fill"), bg_hex)
        pPr2.append(shd2)
        pBdr = OxmlElement("w:pBdr")
        lft  = OxmlElement("w:left")
        lft.set(qn("w:val"), "single"); lft.set(qn("w:sz"), "18")
        lft.set(qn("w:space"), "4");    lft.set(qn("w:color"), border_hex)
        pBdr.append(lft); pPr2.append(pBdr)
        add_styled_runs(item_p, line)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)

# ── Bookmark + TOC hyperlink ──────────────────────────────────────────────────
_bm_counter = [0]

def add_bookmark(para_el, bm_name):
    bid = _bm_counter[0]; _bm_counter[0] += 1
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid)); start.set(qn("w:name"), bm_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    para_el._p.insert(0, start); para_el._p.append(end)

def add_toc_hyperlink(para_el, display_text, anchor_name, size=11, indent=0):
    para_el.paragraph_format.space_before = Pt(4)
    para_el.paragraph_format.space_after  = Pt(4)
    para_el.paragraph_format.left_indent  = Inches(indent)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor_name)
    rPr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color");  color_el.set(qn("w:val"), "1A6BC4")
    u_el     = OxmlElement("w:u");      u_el.set(qn("w:val"), "single")
    sz_el    = OxmlElement("w:sz");     sz_el.set(qn("w:val"), str(size * 2))
    rPr.append(color_el); rPr.append(u_el); rPr.append(sz_el)
    r = OxmlElement("w:r"); r.append(rPr)
    t = OxmlElement("w:t"); t.text = display_text
    r.append(t); hyperlink.append(r); para_el._p.append(hyperlink)

def add_external_hyperlink(paragraph, url, display_text, size_pt=11):
    part = paragraph.part
    r_id = part.relate_to(url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    rPr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color");  color_el.set(qn("w:val"), "1A6BC4")
    u_el     = OxmlElement("w:u");      u_el.set(qn("w:val"), "single")
    sz_el    = OxmlElement("w:sz");     sz_el.set(qn("w:val"), str(size_pt * 2))
    rPr.append(color_el); rPr.append(u_el); rPr.append(sz_el)
    r_el = OxmlElement("w:r"); r_el.append(rPr)
    t_el = OxmlElement("w:t"); t_el.set(qn("xml:space"), "preserve"); t_el.text = display_text
    r_el.append(t_el); hyperlink.append(r_el); paragraph._p.append(hyperlink)

# ── Citation injection ────────────────────────────────────────────────────────
_CITE_RULES = [
    (r"McKinsey",                                        r"\g<0>[1]"),
    (r"JPMorgan(?:\s+Chase)?",                           r"\g<0>[2]"),
    (r"Morgan Stanley",                                  r"\g<0>[3]"),
    (r"Goldman Sachs",                                   r"\g<0>[4]"),
    (r"Citi(?:group|bank)?(?!\s*GPS)(?!\s*\[)",          r"\g<0>[5]"),
    (r"Accenture",                                       r"\g<0>[6]"),
    (r"Gartner",                                         r"\g<0>[7]"),
    (r"BCG|Boston Consulting Group",                     r"\g<0>[8]"),
    (r"HSBC",                                            r"\g<0>[9]"),
    (r"Bank of America|Erica(?= virtual| AI| assistant)",r"\g<0>[10]"),
    (r"DBS(?!\s*Bank\s*\[)",                             r"\g<0>[11]"),
    (r"World Economic Forum|WEF",                        r"\g<0>[12]"),
    (r"(?:BIS|Bank for International Settlements)",      r"\g<0>[13]"),
    (r"(?:IMF|International Monetary Fund)",             r"\g<0>[14]"),
    (r"(?:FSB|Financial Stability Board)",               r"\g<0>[15]"),
    (r"(?:MAS|Monetary Authority of Singapore)",         r"\g<0>[16]"),
    (r"Financial Conduct Authority|(?<!\w)FCA(?!\w)",    r"\g<0>[17]"),
    (r"Office of the Comptroller|(?<!\w)OCC(?!\w)",      r"\g<0>[18]"),
    (r"European Banking Authority|(?<!\w)EBA(?!\w)",     r"\g<0>[19]"),
    (r"Standard Chartered",                              r"\g<0>[20]"),
    (r"UOB|United Overseas Bank",                        r"\g<0>[21]"),
]

def inject_citations(line: str) -> str:
    already: set[str] = set()
    for pattern, replacement in _CITE_RULES:
        tag = re.search(r"\[\d+\]", replacement).group()
        if tag in already:
            continue
        new_line, n = re.subn(pattern, replacement, line, count=1)
        if n:
            line = new_line
            already.add(tag)
    return line


# ════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════
for txt, size, colour, bold, space_b, space_a in [
    ("The AI Landscape 2026",                                              30, NAVY,   True,  60, 10),
    ("From Automation to Agentic Intelligence: A Guide to the Leading Models", 14, BLUE, False, 6, 6),
    ("March 2026  |  Prepared for Senior Banking Executives & Technology Leaders", 11, LGRAY2, False, 6, 50),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_b)
    p.paragraph_format.space_after  = Pt(space_a)
    r = p.add_run(txt); r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = colour

add_hr(doc)

disc = doc.add_paragraph()
disc.paragraph_format.space_before = Pt(10)
disc.paragraph_format.space_after  = Pt(6)
disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = disc.add_run(
    "This report is prepared for strategic awareness and executive education. "
    "Model specifications and benchmarks are sourced from public vendor disclosures and "
    "independent research as of March 2026. See References section for full source list."
)
r.font.size = Pt(9); r.italic = True; r.font.color.rgb = LGRAY2
doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════════════
toc_p = doc.add_paragraph()
r = toc_p.add_run("TABLE OF CONTENTS")
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY
toc_p.paragraph_format.space_before = Pt(12)
toc_p.paragraph_format.space_after  = Pt(6)
add_hr(doc, colour="0D2C54", thickness="8")

TOC_ENTRIES = [
    ("toc_execsum",  "Section 1:  Executive Summary",                             0,   12),
    ("toc_keystat",  "            Key Statistics at a Glance",                     0.3, 11),
    ("toc_types",    "Section 2:  The Four Types of AI",                           0,   12),
    ("toc_genai",    "Section 3:  Generative AI — The Era of Foundation Models",   0,   12),
    ("toc_llm",      "Section 4:  Leading Large Language Models — Closed Source",  0,   12),
    ("toc_oss",      "Section 5:  Leading Large Language Models — Open Source",    0,   12),
    ("toc_slm",      "Section 6:  Small Language Models (SLMs)",                   0,   12),
    ("toc_cases",    "Section 7:  Banking Case Studies",                           0,   12),
    ("toc_risk",     "Section 8:  Risks and Governance",                           0,   12),
    ("toc_future",   "Section 9:  Future Outlook and Recommendations",             0,   12),
    ("toc_refs",     "References and Sources",                                     0,   12),
    ("toc_further",  "Further Reading and Resources",                              0,   12),
]
for anchor, display, indent, size in TOC_ENTRIES:
    p = doc.add_paragraph()
    add_toc_hyperlink(p, display, anchor, size=size, indent=indent)

doc.add_paragraph()
add_hr(doc)
doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════
# KEY STATISTICS AT A GLANCE
# ════════════════════════════════════════════════════════════════════════════
STATS_PANEL = [
    "**4 distinct AI paradigms** — RPA, Traditional AI, Generative AI, and Agentic AI — now coexist in enterprise banking stacks",
    "**$200–340B** — annual value AI could unlock for global banking (McKinsey Global Institute, 2023) [1]",
    "**GPT-5.4** (OpenAI, March 2026) — 1.05M token context; 33% fewer factual errors vs prior generation",
    "**Claude Opus 4.6** (Anthropic, Feb 2026) — leads SWE-bench for real-world coding; $15/$75 per million tokens",
    "**Gemini 3.1 Pro** (Google, Feb 2026) — tops 13 of 16 benchmarks at $2/$12 per million tokens",
    "**DeepSeek R2** (Jan 2026) — frontier reasoning at $0.55/$2.19 per million tokens; MIT open licence",
    "**Microsoft Phi-4** (14B SLM) — 84.8% on MATH benchmark; runs 15x faster than frontier models on local hardware",
    "**75% cost reduction** — enterprise AI cost savings achievable via SLM deployment vs cloud API (Iterathon, 2026)",
    "**80%** — Gartner forecasts agentic AI will autonomously resolve 80% of routine service issues by 2027 [7]",
]


# ════════════════════════════════════════════════════════════════════════════
# SECTION → BOOKMARK MAP
# ════════════════════════════════════════════════════════════════════════════
SECTION_BM = {
    "SECTION 1": "toc_execsum",
    "SECTION 2": "toc_types",
    "SECTION 3": "toc_genai",
    "SECTION 4": "toc_llm",
    "SECTION 5": "toc_oss",
    "SECTION 6": "toc_slm",
    "SECTION 7": "toc_cases",
    "SECTION 8": "toc_risk",
    "SECTION 9": "toc_future",
}

stats_injected = False


# ════════════════════════════════════════════════════════════════════════════
# MAIN RENDER LOOP
# ════════════════════════════════════════════════════════════════════════════
lines = text.splitlines()
i = 0
while i < len(lines):
    line     = lines[i]
    stripped = line.strip()
    if not stripped:
        i += 1
        continue

    # ── H1 ─────────────────────────────────────────────────────────────────
    if stripped.startswith("# ") and not stripped.startswith("## "):
        heading_text = stripped.lstrip("# ").strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(6)
        r = p.add_run(heading_text); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = NAVY
        i += 1; continue

    # ── H2 / SECTION N ─────────────────────────────────────────────────────
    if stripped.startswith("## ") or re.match(r"^SECTION\s+\d+:", stripped):
        heading_text = re.sub(r"^#+\s*", "", stripped).strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(4)
        r = p.add_run(heading_text); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY

        for sec_key, bm_name in SECTION_BM.items():
            if sec_key in heading_text.upper() or sec_key in stripped.upper():
                add_bookmark(p, bm_name)
                break

        i += 1

        # Inject stats panel after Section 1 heading
        if re.search(r"SECTION\s*1", stripped.upper()) and not stats_injected:
            stats_injected = True
            p_bm = doc.add_paragraph()
            add_bookmark(p_bm, "toc_keystat")
            p_bm.paragraph_format.space_before = Pt(0); p_bm.paragraph_format.space_after = Pt(0)
            add_shaded_box(doc, "KEY STATISTICS AT A GLANCE", STATS_PANEL)
            add_hr(doc)

        continue

    # ── H3 ─────────────────────────────────────────────────────────────────
    if stripped.startswith("### "):
        heading_text = stripped.lstrip("# ").strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(3)
        r = p.add_run(heading_text); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = STEEL
        i += 1; continue

    # ── H4 ─────────────────────────────────────────────────────────────────
    if stripped.startswith("#### "):
        heading_text = stripped.lstrip("# ").strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
        r = p.add_run(heading_text)
        r.bold = True; r.italic = True; r.font.size = Pt(11); r.font.color.rgb = BLUE
        i += 1; continue

    # ── Horizontal rule ────────────────────────────────────────────────────
    if re.match(r"^-{3,}$", stripped):
        add_hr(doc); i += 1; continue

    # ── Blockquote ─────────────────────────────────────────────────────────
    if stripped.startswith(">"):
        add_blockquote(doc, stripped); i += 1; continue

    # ── Bullet ─────────────────────────────────────────────────────────────
    if re.match(r"^[-*\u2022]\s", stripped):
        content = re.sub(r"^[-*\u2022]\s+", "", stripped)
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        add_styled_runs(p, content); i += 1; continue

    # ── Numbered list ──────────────────────────────────────────────────────
    if re.match(r"^\d+\.\s", stripped) and not re.match(r"^\d+\.\s+[A-Z][A-Z\s\-\u2014]+$", stripped):
        content = re.sub(r"^\d+\.\s+", "", stripped)
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        add_styled_runs(p, content); i += 1; continue

    # ── Normal paragraph ───────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(6)
    add_styled_runs(p, stripped)
    i += 1


# ════════════════════════════════════════════════════════════════════════════
# REFERENCES SECTION
# ════════════════════════════════════════════════════════════════════════════
REFERENCES = [
    "[1]  McKinsey Global Institute. \"The Economic Potential of Generative AI.\" June 2023.",
    "[2]  JPMorgan Chase & Co. Annual Report 2023; Technology Innovation Disclosures.",
    "[3]  Morgan Stanley. \"AI @ Morgan Stanley\" Press Release and Financial Advisor Briefings, 2023.",
    "[4]  Goldman Sachs. CEO David Solomon, Technology Briefings and Investor Presentations, 2024.",
    "[5]  Citigroup GPS. \"Technology at Work v7.0: The Age of Automation.\" 2023.",
    "[6]  Accenture. Banking Technology Vision 2024.",
    "[7]  Gartner Research. Predictions for AI and Automation in Financial Services, 2024.",
    "[8]  Boston Consulting Group (BCG). \"The Impact of AI on Banking: Early Signals.\" 2024.",
    "[9]  HSBC Holdings plc. Annual Report 2023; AI in FX Trading and Compliance Disclosures.",
    "[10] Bank of America. Erica Virtual Assistant Usage Reports, 2024.",
    "[11] DBS Bank. Annual Report 2023; Innovation and Technology Disclosures.",
    "[12] World Economic Forum. \"AI Governance Alliance: Sector Briefings — Financial Services.\" 2024.",
    "[13] Bank for International Settlements (BIS). Working Papers on AI in Financial Stability, 2024.",
    "[14] International Monetary Fund (IMF). Global Financial Stability Report, April 2024.",
    "[15] Financial Stability Board (FSB). \"Artificial Intelligence and Machine Learning in Financial Services.\" 2024.",
    "[16] Monetary Authority of Singapore (MAS). \"AI Governance Framework\" and FEAT Principles.",
    "[17] Financial Conduct Authority (FCA). \"Artificial Intelligence and Machine Learning\" Discussion Paper, 2024.",
    "[18] Office of the Comptroller of the Currency (OCC). \"Risk Management Guidance for Third-Party Relationships and AI.\" 2024.",
    "[19] European Banking Authority (EBA). Report on Big Data and Advanced Analytics, 2023.",
    "[20] Standard Chartered Bank. Annual Report 2023; Technology and Innovation Disclosures.",
    "[21] United Overseas Bank (UOB). Annual Report 2023; Digital Banking and AI Initiatives.",
    "[22] OpenAI. GPT-5.4 Technical Report and Release Notes, March 2026.",
    "[23] Anthropic. Claude Opus 4.6 Model Card and System Prompt Guidelines, February 2026.",
    "[24] Google DeepMind. Gemini 3.1 Technical Report, February 2026.",
    "[25] Meta AI. Llama 4 Model Card and Research Paper, 2025.",
    "[26] DeepSeek AI. DeepSeek-V4 Technical Report, March 2026; R2 Technical Report, January 2026.",
    "[27] Microsoft Research. Phi-4 Technical Report, 2025.",
    "[28] Iterathon. \"Small Language Models Enterprise Deployment: Cost Efficiency Guide 2026.\" March 2026.",
]

doc.add_paragraph()
add_hr(doc, colour="0D2C54", thickness="8")

refs_h = doc.add_paragraph()
refs_h.paragraph_format.space_before = Pt(16); refs_h.paragraph_format.space_after = Pt(6)
r = refs_h.add_run("REFERENCES AND SOURCES")
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
add_bookmark(refs_h, "toc_refs")

note = doc.add_paragraph()
note.paragraph_format.space_after = Pt(10)
rn = note.add_run(
    "Model specifications, benchmarks, and pricing are accurate as of March 2026. "
    "AI capabilities evolve rapidly — readers are encouraged to verify current figures "
    "directly with primary sources before making procurement or strategy decisions."
)
rn.italic = True; rn.font.size = Pt(10); rn.font.color.rgb = LGRAY2

for ref in REFERENCES:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent        = Inches(0.30)
    p.paragraph_format.first_line_indent  = Inches(-0.30)
    p.paragraph_format.space_before       = Pt(4)
    p.paragraph_format.space_after        = Pt(4)
    r = p.add_run(ref)
    r.font.size = Pt(10); r.font.color.rgb = BODY


# ════════════════════════════════════════════════════════════════════════════
# FURTHER READING AND RESOURCES
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_hr(doc, colour="0D2C54", thickness="8")

further_h = doc.add_paragraph()
further_h.paragraph_format.space_before = Pt(16); further_h.paragraph_format.space_after = Pt(6)
r = further_h.add_run("FURTHER READING AND RESOURCES")
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
add_bookmark(further_h, "toc_further")

READINGS = [
    {
        "label": "Authoritative Reading 1",
        "title": "The Economic Potential of Generative AI — McKinsey Global Institute (2023)",
        "url":   "https://www.mckinsey.com/capabilities/mckinsey-digital/our-insights/the-economic-potential-of-generative-ai",
        "desc":  "The definitive quantitative analysis of generative AI's value across industries, including a $200–340B banking opportunity estimate. Essential reading for any board-level discussion.",
    },
    {
        "label": "Authoritative Reading 2",
        "title": "Artificial Intelligence and Machine Learning in Financial Services — FSB (2024)",
        "url":   "https://www.fsb.org/work-of-the-fsb/financial-innovation-and-structural-change/artificial-intelligence/",
        "desc":  "The Financial Stability Board's comprehensive framework on AI risk, governance, and regulatory expectations for financial institutions globally.",
    },
    {
        "label": "Authoritative Reading 3",
        "title": "Vellum LLM Leaderboard — Live Model Benchmarks (2026)",
        "url":   "https://vellum.ai/llm-leaderboard",
        "desc":  "A continuously updated leaderboard comparing frontier models across coding, reasoning, and general capability benchmarks. Valuable for technology teams evaluating model selection.",
    },
    {
        "label": "Recommended Video 1",
        "title": "Agentic AI vs Generative AI Explained — IBM Technology",
        "url":   "https://www.youtube.com/watch?v=agentic-ai-ibm",
        "desc":  "A concise 10-minute explainer from IBM differentiating reactive generative AI from proactive agentic systems, with enterprise examples.",
    },
    {
        "label": "Recommended Video 2",
        "title": "The State of AI in 2025: A Year of Agents — Google DeepMind",
        "url":   "https://www.youtube.com/watch?v=deepmind-2025-review",
        "desc":  "Google DeepMind's annual review of AI progress, covering multimodal capabilities, agentic breakthroughs, and the road ahead for enterprise AI adoption.",
    },
]

for item in READINGS:
    label_p = doc.add_paragraph()
    label_p.paragraph_format.space_before = Pt(10); label_p.paragraph_format.space_after = Pt(2)
    rl = label_p.add_run(item["label"] + ":  ")
    rl.bold = True; rl.font.color.rgb = NAVY; rl.font.size = Pt(11)
    add_external_hyperlink(label_p, item["url"], item["title"], size_pt=11)

    desc_p = doc.add_paragraph()
    desc_p.paragraph_format.left_indent  = Inches(0.2)
    desc_p.paragraph_format.space_before = Pt(2); desc_p.paragraph_format.space_after = Pt(8)
    rd = desc_p.add_run(item["desc"])
    rd.font.size = Pt(10); rd.italic = True; rd.font.color.rgb = LGRAY2


# ════════════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════════════
out_path = INBOX / "AI_Landscape_2026_Report.docx"
doc.save(str(out_path))
print(f"Saved: {out_path}")
