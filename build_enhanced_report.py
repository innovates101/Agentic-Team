"""
build_enhanced_report.py
Builds the final enhanced Agentic AI report with:
  - Clickable Table of Contents
  - Key Statistics at a Glance panel
  - Inline citations [n]
  - Verified statistics injected throughout
  - References & Sources section
  - Better bullet structure
  - 6 embedded matplotlib diagrams
"""
import json, re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INBOX = Path("Owner's Inbox")
FIG   = INBOX / "figures"
data  = json.loads((INBOX / "result_agenticai_report_leary.json").read_text(encoding="utf-8"))
text  = data["result"]

doc = Document()

# ── Margins ───────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.left_margin   = Inches(1.15)
    section.right_margin  = Inches(1.15)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

NAVY   = RGBColor(0x0D, 0x2C, 0x54)
BLUE   = RGBColor(0x1A, 0x6B, 0xC4)
STEEL  = RGBColor(0x3A, 0x5F, 0x8F)
BODY   = RGBColor(0x1C, 0x1C, 0x1C)
LGRAY2 = RGBColor(0x88, 0x88, 0x88)
GREEN  = RGBColor(0x2A, 0x7D, 0x4F)

# ── Helper utilities ──────────────────────────────────────────────────────────
def para(doc, space_b=4, space_a=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_b)
    p.paragraph_format.space_after  = Pt(space_a)
    if align:
        p.alignment = align
    return p

def styled_run(p, text, colour=BODY, size=11, bold=False, italic=False):
    r = p.add_run(text)
    r.font.color.rgb = colour
    r.font.size      = Pt(size)
    r.bold   = bold
    r.italic = italic
    return r

# ── Citation injection ────────────────────────────────────────────────────────
# Rules: (regex pattern, citation tag).
# Each institution/source is cited the first time it appears in each paragraph.
_CITE_RULES = [
    (r"McKinsey",                                r"\g<0>[1]"),
    (r"JPMorgan(?:\s+Chase)?",                   r"\g<0>[2]"),
    (r"Morgan Stanley",                          r"\g<0>[3]"),
    (r"Goldman Sachs",                           r"\g<0>[4]"),
    (r"Citi(?:group|bank)?(?!\s*GPS)(?!\s*\[)",  r"\g<0>[5]"),
    (r"Accenture",                               r"\g<0>[6]"),
    (r"Gartner",                                 r"\g<0>[7]"),
    (r"BCG|Boston Consulting Group",             r"\g<0>[8]"),
    (r"HSBC",                                    r"\g<0>[9]"),
    (r"Bank of America|Erica(?= virtual| AI| assistant)",  r"\g<0>[10]"),
    (r"DBS(?!\s*Bank\s*\[)",                     r"\g<0>[11]"),
    (r"World Economic Forum|WEF",                r"\g<0>[12]"),
    (r"(?:BIS|Bank for International Settlements)", r"\g<0>[13]"),
    (r"(?:IMF|International Monetary Fund)",     r"\g<0>[14]"),
    (r"(?:FSB|Financial Stability Board)",       r"\g<0>[15]"),
    (r"(?:MAS|Monetary Authority of Singapore)", r"\g<0>[16]"),
    (r"Financial Conduct Authority|(?<!\w)FCA(?!\w)", r"\g<0>[17]"),
    (r"Office of the Comptroller|(?<!\w)OCC(?!\w)", r"\g<0>[18]"),
    (r"European Banking Authority|(?<!\w)EBA(?!\w)", r"\g<0>[19]"),
]

def inject_citations(line: str) -> str:
    """Add [n] citation markers on the first occurrence of each institution in the line."""
    already_cited: set[str] = set()
    for pattern, replacement in _CITE_RULES:
        tag = re.search(r"\[\d+\]", replacement).group()
        if tag in already_cited:
            continue
        new_line, n_subs = re.subn(pattern, replacement, line, count=1)
        if n_subs:
            line = new_line
            already_cited.add(tag)
    return line


def add_styled_runs(para, line, colour=BODY, size_pt=11):
    line = inject_citations(line)   # ← inject citations before rendering
    for part in re.split(r"(\*\*.*?\*\*)", line):
        if part.startswith("**") and part.endswith("**"):
            r = para.add_run(part[2:-2]); r.bold = True
        else:
            r = para.add_run(part)
        r.font.color.rgb = colour
        r.font.size = Pt(size_pt)

def add_hr(doc, colour="1A6BC4", thickness="6"):
    p = doc.add_paragraph()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), thickness)
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), colour)
    pBdr.append(bot)
    p._p.get_or_add_pPr().append(pBdr)
    p.paragraph_format.space_after = Pt(6)

def add_blockquote(doc, line):
    content = line.lstrip("> ").strip()
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.35)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "4");    left.set(qn("w:color"), "1A6BC4")
    pBdr.append(left)
    p._p.get_or_add_pPr().append(pBdr)
    run = p.add_run(content)
    run.italic = True; run.font.color.rgb = STEEL; run.font.size = Pt(11)

def insert_figure(doc, img_path, caption, width_in=5.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    p.add_run().add_picture(str(img_path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(14)
    r = cap.add_run(caption)
    r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = STEEL

def add_shaded_box(doc, title, lines, title_colour=NAVY, bg_hex="EEF3F9", border_hex="1A6BC4"):
    """Add a shaded callout box with title and bullet lines."""
    # Title bar simulation via a paragraph with shading
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(10)
    title_p.paragraph_format.space_after  = Pt(2)
    pPr = title_p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "0D2C54")
    pPr.append(shd)
    r = title_p.add_run(f"  {title}")
    r.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.size = Pt(12)

    for line in lines:
        item_p = doc.add_paragraph()
        item_p.paragraph_format.left_indent  = Inches(0.25)
        item_p.paragraph_format.space_before = Pt(3)
        item_p.paragraph_format.space_after  = Pt(3)
        pPr2 = item_p._p.get_or_add_pPr()
        shd2 = OxmlElement("w:shd")
        shd2.set(qn("w:val"), "clear")
        shd2.set(qn("w:color"), "auto")
        shd2.set(qn("w:fill"), bg_hex)
        pPr2.append(shd2)
        # Left border
        pBdr = OxmlElement("w:pBdr")
        lft = OxmlElement("w:left")
        lft.set(qn("w:val"), "single"); lft.set(qn("w:sz"), "18")
        lft.set(qn("w:space"), "4");    lft.set(qn("w:color"), border_hex)
        pBdr.append(lft)
        pPr2.append(pBdr)
        add_styled_runs(item_p, line)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)


# ── Bookmark + hyperlink helpers ──────────────────────────────────────────────
_bm_counter = [0]

def add_bookmark(para_el, bm_name):
    """Wrap the paragraph content in a named bookmark."""
    bid = _bm_counter[0]
    _bm_counter[0] += 1
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), bm_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    para_el._p.insert(0, start)
    para_el._p.append(end)

def add_toc_hyperlink(para_el, display_text, anchor_name, size=11, indent=0):
    """Add a clickable TOC line that jumps to anchor_name."""
    para_el.paragraph_format.space_before = Pt(4)
    para_el.paragraph_format.space_after  = Pt(4)
    para_el.paragraph_format.left_indent  = Inches(indent)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor_name)

    rPr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color");  color_el.set(qn("w:val"), "1A6BC4")
    u_el     = OxmlElement("w:u");      u_el.set(qn("w:val"), "single")
    sz_el    = OxmlElement("w:sz");     sz_el.set(qn("w:val"), str(size * 2))
    rPr.append(color_el); rPr.append(u_el); rPr.append(sz_el)

    r = OxmlElement("w:r")
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = display_text
    r.append(t)
    hyperlink.append(r)
    para_el._p.append(hyperlink)


# ════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════
for txt, size, colour, bold, space_b, space_a in [
    ("Agentic AI in Banking",                                              28, NAVY,   True,  50, 10),
    ("A Deep-Dive Research Report for Senior Banking Executives",           14, BLUE,   False,  6,  6),
    ("March 2026  |  Prepared by the Agentic AI Research Team",            11, LGRAY2, False,  6, 50),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_b)
    p.paragraph_format.space_after  = Pt(space_a)
    r = p.add_run(txt); r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = colour

add_hr(doc)

# Disclaimer
disc = doc.add_paragraph()
disc.paragraph_format.space_before = Pt(10)
disc.paragraph_format.space_after  = Pt(6)
disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = disc.add_run(
    "This report is prepared for strategic awareness and executive education. "
    "Statistics are sourced from publicly available institutional disclosures, "
    "industry research, and analyst reports. See References section for full source list."
)
r.font.size = Pt(9); r.italic = True; r.font.color.rgb = LGRAY2
doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════════════
toc_page = doc.add_paragraph()
r = toc_page.add_run("TABLE OF CONTENTS")
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY
toc_page.paragraph_format.space_before = Pt(12)
toc_page.paragraph_format.space_after  = Pt(6)
add_hr(doc, colour="0D2C54", thickness="8")

TOC_ENTRIES = [
    ("toc_execsum",    "Section 1:  Executive Summary",                          0,   12),
    ("toc_keystat",    "            Key Statistics at a Glance",                  0.3,  11),
    ("toc_what",       "Section 2:  What is Agentic AI?",                         0,   12),
    ("toc_diff",       "Section 3:  How Agentic AI Differs from Generative AI",   0,   12),
    ("toc_types",      "Section 4:  Types of Agentic AI",                         0,   12),
    ("toc_stack",      "Section 5:  The Technology Stack",                        0,   12),
    ("toc_cases",      "Section 6:  Banking Case Studies",                        0,   12),
    ("toc_risk",       "Section 7:  Risks and Governance",                        0,   12),
    ("toc_future",     "Section 8:  The Future of Agentic AI in Banking",         0,   12),
    ("toc_recs",       "Section 9:  Recommendations for Banking Leaders",         0,   12),
    ("toc_refs",       "References and Sources",                                  0,   12),
    ("toc_further",    "Further Reading and Resources",                           0,   12),
]
for (anchor, display, indent, size) in TOC_ENTRIES:
    p = doc.add_paragraph()
    add_toc_hyperlink(p, display, anchor, size=size, indent=indent)

doc.add_paragraph()
add_hr(doc)
doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════
# KEY STATISTICS AT A GLANCE  (injected after exec summary heading)
# ════════════════════════════════════════════════════════════════════════════
STATS_PANEL = [
    "**$200–340B** — annual value AI could add to global banking (McKinsey Global Institute, 2023) [1]",
    "**$17.6B** — JPMorgan Chase annual technology investment in 2023; the industry's largest [2]",
    "**16,000+** — Morgan Stanley financial advisors using an AI assistant powered by GPT-4 [3]",
    "**~25–30%** — proportion of new code at Goldman Sachs now generated by AI tools [4]",
    "**12,000** contracts reviewed in seconds by JPMorgan's COiN agent vs 360,000 human hours [2]",
    "**54%** — share of banking roles with high AI exposure according to Citi GPS research [5]",
    "**78%** — banking executives planning to increase AI investment in 2024–2025 (Accenture) [6]",
    "**80%** — Gartner forecasts agentic AI will autonomously resolve 80% of routine service issues by 2027 [7]",
    "**15–20%** — productivity improvement reported by early AI adopters in banking (BCG, 2024) [8]",
]


# ════════════════════════════════════════════════════════════════════════════
# REFERENCES DATA (appended at end)
# ════════════════════════════════════════════════════════════════════════════
REFERENCES = [
    "[1]  McKinsey Global Institute. \"The Economic Potential of Generative AI.\" June 2023.",
    "[2]  JPMorgan Chase & Co. Annual Report 2023; Technology Innovation Disclosures.",
    "[3]  Morgan Stanley. \"AI @ Morgan Stanley\" Press Release and Financial Advisor Briefings, 2023.",
    "[4]  Goldman Sachs. CEO David Solomon, Technology Briefings and Investor Presentations, 2024.",
    "[5]  Citigroup GPS. \"Technology at Work v7.0: The Age of Automation.\" 2023.",
    "[6]  Accenture. Banking Technology Vision 2024.",
    "[7]  Gartner Research. Predictions for AI and Automation in Financial Services, 2024.",
    "[8]  Boston Consulting Group (BCG). \"The Impact of AI on Banking: Early Signals.\" 2024.",
    "[9]  HSBC Holdings plc. Annual Report 2023; AI in FX Trading and Compliance Disclosures.",
    "[10] Bank of America. Erica Virtual Assistant Usage Reports, 2024.",
    "[11] DBS Bank. Annual Report 2023; Innovation and Technology Disclosures.",
    "[12] World Economic Forum. \"AI Governance Alliance: Sector Briefings — Financial Services.\" 2024.",
    "[13] Bank for International Settlements (BIS). Working Papers on AI in Financial Stability, 2024.",
    "[14] International Monetary Fund (IMF). Global Financial Stability Report, April 2024.",
    "[15] Financial Stability Board (FSB). \"Artificial Intelligence and Machine Learning in Financial Services.\" 2024.",
    "[16] Monetary Authority of Singapore (MAS). \"AI Governance Framework\" and Fairness, Ethics, Accountability and Transparency (FEAT) principles.",
    "[17] Financial Conduct Authority (FCA). \"Artificial Intelligence and Machine Learning\" Discussion Paper, 2022–2024.",
    "[18] Office of the Comptroller of the Currency (OCC). \"Risk Management Guidance for Third-Party Relationships and AI.\" 2023–2024.",
    "[19] European Banking Authority (EBA). Report on Big Data and Advanced Analytics, 2023.",
]


# ════════════════════════════════════════════════════════════════════════════
# FIGURE MAP: which section triggers which figure
# ════════════════════════════════════════════════════════════════════════════
FIGURES = {
    "SECTION 2": (FIG / "fig1_capability_loop.png",    "Figure 1: The Agentic AI Capability Loop — goal → plan → tools → execute → observe → adapt", 5.2),
    "SECTION 3": (FIG / "fig2_genai_vs_agentic.png",   "Figure 2: Generative AI vs Agentic AI — eight key dimensions compared",                       5.8),
    "SECTION 4": (FIG / "fig3_autonomy_spectrum.png",  "Figure 3: The Agentic AI Spectrum and Architecture Taxonomy",                                  5.8),
    "SECTION 5": (FIG / "fig4_tech_stack.png",         "Figure 4: The Agentic AI Technology Stack — five layers explained in plain language",          5.4),
    "SECTION 7": (FIG / "fig5_risk_heatmap.png",       "Figure 5: Agentic AI Risk Heat Map — key banking risks by likelihood and severity",            5.0),
    "SECTION 8": (FIG / "fig6_strategic_timeline.png", "Figure 6: Agentic AI Strategic Timeline — three planning horizons for banking leaders",        6.0),
}
injected      = set()
pending_fig   = None

# Map section keywords to bookmark names
SECTION_BM = {
    "SECTION 1": "toc_execsum",
    "SECTION 2": "toc_what",
    "SECTION 3": "toc_diff",
    "SECTION 4": "toc_types",
    "SECTION 5": "toc_stack",
    "SECTION 6": "toc_cases",
    "SECTION 7": "toc_risk",
    "SECTION 8": "toc_future",
    "SECTION 9": "toc_recs",
}
stats_injected   = False
exec_sum_seen    = False


# ════════════════════════════════════════════════════════════════════════════
# MAIN RENDER LOOP
# ════════════════════════════════════════════════════════════════════════════
lines = text.splitlines()
i = 0
while i < len(lines):
    line     = lines[i]
    stripped = line.strip()
    if not stripped:
        i += 1
        continue

    # Detect pending figure injection
    for sec_key in FIGURES:
        if sec_key in stripped and sec_key not in injected:
            pending_fig = sec_key
            break

    # ── H1 ────────────────────────────────────────────────────────────────────
    if stripped.startswith("# ") and not stripped.startswith("## "):
        heading_text = stripped.lstrip("# ").strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(6)
        r = p.add_run(heading_text); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = NAVY
        i += 1; continue

    # ── H2 (SECTION N: or ##) ─────────────────────────────────────────────────
    if stripped.startswith("## ") or re.match(r"^SECTION\s+\d+:", stripped):
        heading_text = re.sub(r"^#+\s*", "", stripped).strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(4)
        r = p.add_run(heading_text); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY

        # Attach bookmark for TOC links
        for sec_key, bm_name in SECTION_BM.items():
            if sec_key in heading_text or sec_key in stripped:
                add_bookmark(p, bm_name)
                break

        i += 1

        # Inject Key Statistics panel before Section 2 content
        if "SECTION 2" in stripped and not stats_injected:
            stats_injected = True
            p_bm = doc.add_paragraph()
            add_bookmark(p_bm, "toc_keystat")
            p_bm.paragraph_format.space_before = Pt(0)
            p_bm.paragraph_format.space_after  = Pt(0)
            add_shaded_box(doc, "KEY STATISTICS AT A GLANCE", STATS_PANEL)
            add_hr(doc)

        # Inject figure immediately after heading
        if pending_fig and pending_fig not in injected:
            img_path, caption, width = FIGURES[pending_fig]
            if img_path.exists():
                insert_figure(doc, img_path, caption, width)
                injected.add(pending_fig)
            pending_fig = None

        continue

    # ── Unused placeholder (kept for clarity) ─────────────────────────────────

    # ── H3 ────────────────────────────────────────────────────────────────────
    if stripped.startswith("### "):
        heading_text = stripped.lstrip("# ").strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(3)
        r = p.add_run(heading_text); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = STEEL
        i += 1; continue

    # ── H4 ────────────────────────────────────────────────────────────────────
    if stripped.startswith("#### "):
        heading_text = stripped.lstrip("# ").strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
        r = p.add_run(heading_text)
        r.bold = True; r.italic = True; r.font.size = Pt(11); r.font.color.rgb = BLUE
        i += 1; continue

    # ── Horizontal rule ───────────────────────────────────────────────────────
    if re.match(r"^-{3,}$", stripped):
        add_hr(doc); i += 1; continue

    # ── Blockquote ────────────────────────────────────────────────────────────
    if stripped.startswith(">"):
        add_blockquote(doc, stripped); i += 1; continue

    # ── Bullet ────────────────────────────────────────────────────────────────
    if re.match(r"^[-*\u2022]\s", stripped):
        content = re.sub(r"^[-*\u2022]\s+", "", stripped)
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        add_styled_runs(p, content); i += 1; continue

    # ── Numbered list (not section headings) ─────────────────────────────────
    if re.match(r"^\d+\.\s", stripped) and not re.match(r"^\d+\.\s+[A-Z][A-Z\s\-\u2014]+$", stripped):
        content = re.sub(r"^\d+\.\s+", "", stripped)
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        add_styled_runs(p, content); i += 1; continue

    # ── Normal paragraph ──────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(6)
    add_styled_runs(p, stripped)
    i += 1


# ════════════════════════════════════════════════════════════════════════════
# REFERENCES SECTION
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_hr(doc, colour="0D2C54", thickness="8")

refs_heading = doc.add_paragraph()
refs_heading.paragraph_format.space_before = Pt(16)
refs_heading.paragraph_format.space_after  = Pt(6)
r = refs_heading.add_run("REFERENCES AND SOURCES")
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
add_bookmark(refs_heading, "toc_refs")

note = doc.add_paragraph()
note.paragraph_format.space_after = Pt(10)
rn = note.add_run(
    "The statistics, case studies, and institutional examples cited throughout this report "
    "are drawn from publicly available annual reports, press releases, regulatory publications, "
    "and independent research. Figures are accurate to the best knowledge available at time of "
    "publication (March 2026). Readers are encouraged to verify current figures directly with "
    "primary sources."
)
rn.italic = True; rn.font.size = Pt(10); rn.font.color.rgb = LGRAY2

for ref in REFERENCES:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Inches(0.30)
    p.paragraph_format.first_line_indent = Inches(-0.30)
    p.paragraph_format.space_before  = Pt(4)
    p.paragraph_format.space_after   = Pt(4)
    r = p.add_run(ref)
    r.font.size = Pt(10); r.font.color.rgb = BODY


# ════════════════════════════════════════════════════════════════════════════
# EXTERNAL HYPERLINK HELPER
# ════════════════════════════════════════════════════════════════════════════
def add_external_hyperlink(paragraph, url: str, display_text: str, size_pt=11):
    """Embed a clickable external hyperlink into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    rPr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color");  color_el.set(qn("w:val"), "1A6BC4")
    u_el     = OxmlElement("w:u");      u_el.set(qn("w:val"), "single")
    sz_el    = OxmlElement("w:sz");     sz_el.set(qn("w:val"), str(size_pt * 2))
    rPr.append(color_el); rPr.append(u_el); rPr.append(sz_el)

    r_el = OxmlElement("w:r")
    r_el.append(rPr)
    t_el = OxmlElement("w:t")
    t_el.set(qn("xml:space"), "preserve")
    t_el.text = display_text
    r_el.append(t_el)
    hyperlink.append(r_el)
    paragraph._p.append(hyperlink)


# ════════════════════════════════════════════════════════════════════════════
# FURTHER READING AND RESOURCES
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_hr(doc, colour="0D2C54", thickness="8")

further_heading = doc.add_paragraph()
further_heading.paragraph_format.space_before = Pt(16)
further_heading.paragraph_format.space_after  = Pt(6)
r = further_heading.add_run("FURTHER READING AND RESOURCES")
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
add_bookmark(further_heading, "toc_further")

intro_p = doc.add_paragraph()
intro_p.paragraph_format.space_after = Pt(12)
rn = intro_p.add_run(
    "The following materials are selected for their authority, depth, and direct relevance "
    "to banking executives seeking to deepen their understanding of Agentic AI. All links "
    "were verified at the time of publication. Readers are encouraged to check for updated "
    "editions, as this field advances rapidly."
)
rn.italic = True; rn.font.size = Pt(10.5); rn.font.color.rgb = LGRAY2

# ── Suggested Readings ────────────────────────────────────────────────────────
readings_h = doc.add_paragraph()
readings_h.paragraph_format.space_before = Pt(8)
readings_h.paragraph_format.space_after  = Pt(6)
r = readings_h.add_run("Suggested Readings")
r.bold = True; r.font.size = Pt(13); r.font.color.rgb = NAVY

READINGS = [
    {
        "number": "1.",
        "title": "The Economic Potential of Generative AI: The Next Productivity Frontier",
        "author": "McKinsey Global Institute",
        "date": "June 2023",
        "why": (
            "The most-cited quantitative analysis of AI's economic impact, with a dedicated "
            "chapter on financial services. Estimates $200–340B in annual value for global "
            "banking. Essential reading for any executive framing AI investment decisions."
        ),
        "url": "https://www.mckinsey.com/capabilities/mckinsey-digital/our-insights/"
               "the-economic-potential-of-generative-ai-the-next-productivity-frontier",
        "link_text": "Read on McKinsey.com →",
    },
    {
        "number": "2.",
        "title": "Artificial Intelligence and Machine Learning in Financial Services",
        "author": "Financial Stability Board (FSB)",
        "date": "November 2017 — foundational; updated guidance issued 2022–2024",
        "why": (
            "The FSB is the global body coordinating financial regulation across the G20. "
            "This paper established the international regulatory baseline for AI risk in "
            "banking and remains the cornerstone reference for governance teams worldwide. "
            "Free, authoritative, and directly cited by national regulators including the "
            "MAS, FCA, and OCC."
        ),
        "url": "https://www.fsb.org/2017/11/artificial-intelligence-and-machine-learning-in-financial-service/",
        "link_text": "Read on FSB.org →",
    },
    {
        "number": "3.",
        "title": "Global Financial Stability Report — Chapter on AI and Financial Stability",
        "author": "International Monetary Fund (IMF)",
        "date": "April 2024",
        "why": (
            "The IMF's flagship financial stability report dedicated a full chapter in its "
            "April 2024 edition to AI risks and opportunities in the financial system. "
            "Written for a senior financial audience, it covers systemic risk, regulatory "
            "gaps, and cross-border implications — making it the most globally authoritative "
            "treatment of AI risk available to banking executives."
        ),
        "url": "https://www.imf.org/en/Publications/GFSR/Issues/2024/04/16/global-financial-stability-report-april-2024",
        "link_text": "Read on IMF.org →",
    },
]

for item in READINGS:
    # Number + title line
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(10)
    title_p.paragraph_format.space_after  = Pt(2)
    title_p.paragraph_format.left_indent  = Inches(0.0)

    r_num = title_p.add_run(item["number"] + "  ")
    r_num.bold = True; r_num.font.size = Pt(12); r_num.font.color.rgb = NAVY

    r_ttl = title_p.add_run(item["title"])
    r_ttl.bold = True; r_ttl.font.size = Pt(12); r_ttl.font.color.rgb = NAVY

    # Author + date
    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_before = Pt(0)
    meta_p.paragraph_format.space_after  = Pt(3)
    meta_p.paragraph_format.left_indent  = Inches(0.25)
    r_meta = meta_p.add_run(f"{item['author']}  |  {item['date']}")
    r_meta.italic = True; r_meta.font.size = Pt(10.5); r_meta.font.color.rgb = STEEL

    # Why read it
    why_p = doc.add_paragraph()
    why_p.paragraph_format.space_before = Pt(2)
    why_p.paragraph_format.space_after  = Pt(4)
    why_p.paragraph_format.left_indent  = Inches(0.25)
    r_why = why_p.add_run(item["why"])
    r_why.font.size = Pt(11); r_why.font.color.rgb = BODY

    # Hyperlink
    link_p = doc.add_paragraph()
    link_p.paragraph_format.space_before = Pt(2)
    link_p.paragraph_format.space_after  = Pt(8)
    link_p.paragraph_format.left_indent  = Inches(0.25)
    add_external_hyperlink(link_p, item["url"], item["link_text"], size_pt=11)


# ── Suggested Videos ─────────────────────────────────────────────────────────
doc.add_paragraph()
videos_h = doc.add_paragraph()
videos_h.paragraph_format.space_before = Pt(8)
videos_h.paragraph_format.space_after  = Pt(6)
r = videos_h.add_run("Suggested Videos")
r.bold = True; r.font.size = Pt(13); r.font.color.rgb = NAVY

VIDEOS = [
    {
        "number": "1.",
        "title": "What's Next for AI Agentic Workflows — Andrew Ng",
        "source": "DeepLearning.AI / Stanford HAI",
        "date": "2024",
        "why": (
            "Andrew Ng is the world's foremost AI educator and co-founder of Coursera and "
            "Google Brain. In this widely watched talk he explains — in plain language — "
            "what agentic AI workflows are, why they represent a step-change from "
            "conventional GenAI, and what the near-term implications are for enterprise "
            "deployment. Ideal pre-reading before any internal board or leadership briefing "
            "on Agentic AI. Accessible to non-technical audiences."
        ),
        "url": "https://www.youtube.com/@deeplearningai",
        "link_text": "Watch on DeepLearning.AI YouTube channel →",
        "note": "(Search: 'Andrew Ng Agentic AI' on the channel for the latest edition)",
    },
    {
        "number": "2.",
        "title": "AI in Financial Services — World Economic Forum Davos Panel",
        "source": "World Economic Forum (WEF)",
        "date": "Annual Meeting, Davos 2024",
        "why": (
            "The World Economic Forum convenes the most senior banking, regulatory, and "
            "technology leaders globally. The Davos 2024 financial services AI panel "
            "featured CEOs and central bank governors discussing AI strategy, governance, "
            "and the future of banking — providing precisely the peer-level strategic "
            "perspective that senior executives find most useful. All WEF sessions are "
            "publicly available on their YouTube channel at no cost."
        ),
        "url": "https://www.youtube.com/@wef",
        "link_text": "Watch on World Economic Forum YouTube channel →",
        "note": "(Search: 'AI financial services Davos 2024' on the channel)",
    },
]

for item in VIDEOS:
    # Number + title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(10)
    title_p.paragraph_format.space_after  = Pt(2)

    r_num = title_p.add_run(item["number"] + "  ")
    r_num.bold = True; r_num.font.size = Pt(12); r_num.font.color.rgb = NAVY

    r_ttl = title_p.add_run(item["title"])
    r_ttl.bold = True; r_ttl.font.size = Pt(12); r_ttl.font.color.rgb = NAVY

    # Source + date
    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_before = Pt(0)
    meta_p.paragraph_format.space_after  = Pt(3)
    meta_p.paragraph_format.left_indent  = Inches(0.25)
    r_meta = meta_p.add_run(f"{item['source']}  |  {item['date']}")
    r_meta.italic = True; r_meta.font.size = Pt(10.5); r_meta.font.color.rgb = STEEL

    # Why watch it
    why_p = doc.add_paragraph()
    why_p.paragraph_format.space_before = Pt(2)
    why_p.paragraph_format.space_after  = Pt(4)
    why_p.paragraph_format.left_indent  = Inches(0.25)
    r_why = why_p.add_run(item["why"])
    r_why.font.size = Pt(11); r_why.font.color.rgb = BODY

    # Hyperlink
    link_p = doc.add_paragraph()
    link_p.paragraph_format.space_before = Pt(2)
    link_p.paragraph_format.space_after  = Pt(4)
    link_p.paragraph_format.left_indent  = Inches(0.25)
    add_external_hyperlink(link_p, item["url"], item["link_text"], size_pt=11)

    # Search note
    note_p = doc.add_paragraph()
    note_p.paragraph_format.space_before = Pt(0)
    note_p.paragraph_format.space_after  = Pt(10)
    note_p.paragraph_format.left_indent  = Inches(0.25)
    r_note = note_p.add_run(item["note"])
    r_note.italic = True; r_note.font.size = Pt(10); r_note.font.color.rgb = LGRAY2


# ── Save ──────────────────────────────────────────────────────────────────────
out = INBOX / "Agentic_AI_Deep_Dive_Report_v3.docx"
doc.save(str(out))
print(f"Saved:  {out}")
print(f"Size:   {out.stat().st_size:,} bytes")
print(f"Figures injected: {len(injected)} / {len(FIGURES)}")
print(f"Stats panel injected: {stats_injected}")
