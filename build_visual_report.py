"""
build_visual_report.py
Rebuilds the Agentic AI report Word document with 6 embedded matplotlib diagrams.
"""
import json, re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INBOX  = Path("Owner's Inbox")
FIG    = INBOX / "figures"
data   = json.loads((INBOX / "result_agenticai_report_leary.json").read_text(encoding="utf-8"))
text   = data["result"]

doc = Document()

# ── Margins ───────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.left_margin   = Inches(1.15)
    section.right_margin  = Inches(1.15)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

NAVY   = RGBColor(0x0D, 0x2C, 0x54)
BLUE   = RGBColor(0x1A, 0x6B, 0xC4)
STEEL  = RGBColor(0x3A, 0x5F, 0x8F)
BODY   = RGBColor(0x1C, 0x1C, 0x1C)
LGRAY2 = RGBColor(0x88, 0x88, 0x88)


# ── Helpers ───────────────────────────────────────────────────────────────────
def add_styled_runs(para, line, colour=BODY, size_pt=11):
    for part in re.split(r"(\*\*.*?\*\*)", line):
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            run = para.add_run(part)
        run.font.color.rgb = colour
        run.font.size = Pt(size_pt)


def add_hr(doc):
    p = doc.add_paragraph()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1A6BC4")
    pBdr.append(bot)
    p._p.get_or_add_pPr().append(pBdr)
    p.paragraph_format.space_after = Pt(6)


def add_blockquote(doc, line):
    content = line.lstrip("> ").strip()
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.35)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), "1A6BC4")
    pBdr.append(left)
    p._p.get_or_add_pPr().append(pBdr)
    run = p.add_run(content)
    run.italic = True
    run.font.color.rgb = STEEL
    run.font.size = Pt(11)


def insert_figure(doc, img_path, caption, width_in=5.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    p.add_run().add_picture(str(img_path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(14)
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = STEEL


# ── Cover page ────────────────────────────────────────────────────────────────
for txt, size, colour, bold, space_b, space_a in [
    ("Agentic AI in Banking", 28, NAVY, True, 40, 8),
    ("A Deep-Dive Research Report for Senior Banking Executives", 14, BLUE, False, 6, 6),
    ("March 2026  |  Prepared by the Agentic AI Research Team", 11, LGRAY2, False, 6, 40),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_b)
    p.paragraph_format.space_after  = Pt(space_a)
    r = p.add_run(txt)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = colour

add_hr(doc)
doc.add_paragraph()

# ── Figure map: section keyword → (image, caption, width_inches) ──────────────
FIGURES = {
    "SECTION 2": (FIG / "fig1_capability_loop.png",   "Figure 1: The Agentic AI Capability Loop — how an agent pursues a goal autonomously",                 5.0),
    "SECTION 3": (FIG / "fig2_genai_vs_agentic.png",  "Figure 2: Generative AI vs Agentic AI — eight dimensions compared",                                   5.5),
    "SECTION 4": (FIG / "fig3_autonomy_spectrum.png",  "Figure 3: The Agentic AI Spectrum and Taxonomy — from human-led to fully autonomous",                 5.5),
    "SECTION 5": (FIG / "fig4_tech_stack.png",         "Figure 4: The Agentic AI Technology Stack — five layers explained",                                   5.2),
    "SECTION 7": (FIG / "fig5_risk_heatmap.png",       "Figure 5: Agentic AI Risk Heat Map — key risks by likelihood and severity",                           4.8),
    "SECTION 8": (FIG / "fig6_strategic_timeline.png", "Figure 6: Agentic AI Strategic Timeline — three horizons for banking leaders",                        5.8),
}
injected = set()
pending_fig = None   # figure to inject after the next heading render


# ── Parse and render ──────────────────────────────────────────────────────────
lines = text.splitlines()
i = 0
while i < len(lines):
    line    = lines[i]
    stripped = line.strip()

    if not stripped:
        i += 1
        continue

    # Check if this line signals a section for which we have a figure
    for sec_key in FIGURES:
        if sec_key in stripped and sec_key not in injected:
            pending_fig = sec_key
            break

    # H1
    if stripped.startswith("# ") and not stripped.startswith("## "):
        heading_text = stripped.lstrip("# ").strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
        r = p.add_run(heading_text)
        r.bold = True; r.font.size = Pt(18); r.font.color.rgb = NAVY
        i += 1
        continue

    # H2 (## or SECTION N: style)
    if stripped.startswith("## ") or re.match(r"^SECTION\s+\d+:", stripped):
        heading_text = re.sub(r"^#+\s*", "", stripped).strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(heading_text)
        r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
        i += 1
        # Inject figure immediately after section heading
        if pending_fig and pending_fig not in injected:
            img_path, caption, width = FIGURES[pending_fig]
            if img_path.exists():
                insert_figure(doc, img_path, caption, width)
                injected.add(pending_fig)
            pending_fig = None
        continue

    # H3
    if stripped.startswith("### "):
        heading_text = stripped.lstrip("# ").strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(heading_text)
        r.bold = True; r.font.size = Pt(12); r.font.color.rgb = STEEL
        i += 1
        continue

    # H4
    if stripped.startswith("#### "):
        heading_text = stripped.lstrip("# ").strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(heading_text)
        r.bold = True; r.italic = True
        r.font.size = Pt(11); r.font.color.rgb = BLUE
        i += 1
        continue

    # Horizontal rule
    if re.match(r"^-{3,}$", stripped):
        add_hr(doc)
        i += 1
        continue

    # Blockquote
    if stripped.startswith(">"):
        add_blockquote(doc, stripped)
        i += 1
        continue

    # Bullet
    if re.match(r"^[-*\u2022]\s", stripped):
        content = re.sub(r"^[-*\u2022]\s+", "", stripped)
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        add_styled_runs(p, content)
        i += 1
        continue

    # Numbered list (but not SECTION headings)
    if re.match(r"^\d+\.\s", stripped) and not re.match(r"^\d+\.\s+[A-Z][A-Z\s\-\u2014]+$", stripped):
        content = re.sub(r"^\d+\.\s+", "", stripped)
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        add_styled_runs(p, content)
        i += 1
        continue

    # Normal paragraph
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    add_styled_runs(p, stripped)
    i += 1


# ── Save ──────────────────────────────────────────────────────────────────────
out = INBOX / "Agentic_AI_Deep_Dive_Report_Visual.docx"
doc.save(str(out))
print(f"Saved: {out}  ({out.stat().st_size:,} bytes)")
print(f"Figures injected: {len(injected)} / {len(FIGURES)}")
print(f"Injected sections: {sorted(injected)}")
